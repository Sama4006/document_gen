"""
Create multiple sample Word templates with various styles and placeholder patterns
Run: python create_templates.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_report_template():
    """Professional business report template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('{{REPORT_TITLE}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = doc.add_paragraph()
    meta.add_run('Prepared by: ').bold = True
    meta.add_run('{{AUTHOR}}\n')
    meta.add_run('Date: ').bold = True
    meta.add_run('{{DATE}}\n')
    meta.add_run('Department: ').bold = True
    meta.add_run('{{DEPARTMENT}}')
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('{{EXECUTIVE_SUMMARY}}')
    
    # Key Findings
    doc.add_heading('Key Findings', level=1)
    doc.add_paragraph('{{KEY_FINDINGS}}')
    
    # Analysis
    doc.add_heading('Detailed Analysis', level=1)
    doc.add_paragraph('{{ANALYSIS}}')
    
    # Recommendations
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('{{RECOMMENDATIONS}}')
    
    # Metrics Table
    doc.add_heading('Performance Metrics', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    metrics = [
        ('Progress', '{{PROGRESS}}'),
        ('Status', '{{STATUS}}'),
        ('Impact', '{{IMPACT}}')
    ]
    
    for i, (metric, value) in enumerate(metrics, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('{{CONCLUSION}}')
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/report_template.docx')
    print('✓ Report template created')


def create_proposal_template():
    """Business proposal template"""
    doc = Document()
    
    # Header
    header_para = doc.add_heading('Business Proposal', 0)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('{{PROJECT_NAME}}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    
    # Client Info
    doc.add_heading('Client Information', level=1)
    para = doc.add_paragraph()
    para.add_run('Client: ').bold = True
    para.add_run('{{CLIENT_NAME}}\n')
    para.add_run('Contact: ').bold = True
    para.add_run('{{CLIENT_CONTACT}}\n')
    para.add_run('Date Submitted: ').bold = True
    para.add_run('{{SUBMISSION_DATE}}')
    
    # Proposal Overview
    doc.add_heading('Proposal Overview', level=1)
    doc.add_paragraph('{{PROPOSAL_OVERVIEW}}')
    
    # Scope of Work
    doc.add_heading('Scope of Work', level=1)
    doc.add_paragraph('{{SCOPE_OF_WORK}}')
    
    # Deliverables
    doc.add_heading('Deliverables', level=1)
    doc.add_paragraph('{{DELIVERABLES}}')
    
    # Timeline
    doc.add_heading('Timeline & Milestones', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    header = table.rows[0].cells
    header[0].text = 'Milestone'
    header[1].text = 'Target Date'
    
    milestones = [
        ('Phase 1: Planning', '{{PHASE1_DATE}}'),
        ('Phase 2: Execution', '{{PHASE2_DATE}}'),
        ('Phase 3: Delivery', '{{PHASE3_DATE}}')
    ]
    
    for i, (milestone, date) in enumerate(milestones, 1):
        table.rows[i].cells[0].text = milestone
        table.rows[i].cells[1].text = date
    
    # Investment
    doc.add_heading('Investment & Terms', level=1)
    para = doc.add_paragraph()
    para.add_run('Total Investment: ').bold = True
    para.add_run('{{INVESTMENT}}\n')
    para.add_run('Payment Terms: ').bold = True
    para.add_run('{{PAYMENT_TERMS}}')
    
    # Next Steps
    doc.add_heading('Next Steps', level=1)
    doc.add_paragraph('{{NEXT_STEPS}}')
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/proposal_template.docx')
    print('✓ Proposal template created')


def create_resume_template():
    """Professional resume template"""
    doc = Document()
    
    # Header
    name = doc.add_heading('{{FULL_NAME}}', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph('{{EMAIL}} | {{PHONE}} | {{LOCATION}}')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(10)
    
    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('{{PROFESSIONAL_SUMMARY}}')
    
    # Experience
    doc.add_heading('Professional Experience', level=1)
    doc.add_paragraph('{{WORK_EXPERIENCE}}')
    
    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph('{{EDUCATION}}')
    
    # Skills
    doc.add_heading('Core Competencies', level=1)
    doc.add_paragraph('{{SKILLS}}')
    
    # Certifications
    doc.add_heading('Certifications', level=1)
    doc.add_paragraph('{{CERTIFICATIONS}}')
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/resume_template.docx')
    print('✓ Resume template created')


def create_email_template():
    """Professional email template"""
    doc = Document()
    
    # Subject
    subject = doc.add_paragraph()
    subject.add_run('Subject: ').bold = True
    subject.add_run('{{EMAIL_SUBJECT}}')
    
    doc.add_paragraph()  # Spacing
    
    # Greeting
    doc.add_paragraph('{{GREETING}}')
    
    # Body
    doc.add_paragraph('{{EMAIL_BODY}}')
    
    # Closing
    doc.add_paragraph()  # Spacing
    doc.add_paragraph('{{CLOSING}}')
    
    # Signature
    sig = doc.add_paragraph()
    sig.add_run('\n').bold = True
    sig.add_run('{{SENDER_NAME}}\n')
    sig.add_run('{{SENDER_TITLE}}\n')
    sig.add_run('{{SENDER_CONTACT}}')
    sig.runs[0].font.size = Pt(10)
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/email_template.docx')
    print('✓ Email template created')


def create_contract_template():
    """Legal contract template"""
    doc = Document()
    
    # Title
    title = doc.add_heading('{{CONTRACT_TITLE}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Effective Date
    doc.add_paragraph(f'Effective Date: {{{{EFFECTIVE_DATE}}}}')
    
    # Parties
    doc.add_heading('Parties', level=1)
    para = doc.add_paragraph()
    para.add_run('This Agreement is entered into between:\n\n')
    para.add_run('Party A: ').bold = True
    para.add_run('{{PARTY_A_NAME}}\n')
    para.add_run('Address: ').bold = True
    para.add_run('{{PARTY_A_ADDRESS}}\n\n')
    para.add_run('Party B: ').bold = True
    para.add_run('{{PARTY_B_NAME}}\n')
    para.add_run('Address: ').bold = True
    para.add_run('{{PARTY_B_ADDRESS}}')
    
    # Terms & Conditions
    doc.add_heading('Terms & Conditions', level=1)
    doc.add_paragraph('{{TERMS_CONDITIONS}}')
    
    # Payment Terms
    doc.add_heading('Payment Terms', level=1)
    doc.add_paragraph('{{PAYMENT_TERMS}}')
    
    # Confidentiality
    doc.add_heading('Confidentiality', level=1)
    doc.add_paragraph('{{CONFIDENTIALITY}}')
    
    # Termination
    doc.add_heading('Termination', level=1)
    doc.add_paragraph('{{TERMINATION}}')
    
    # Signatures
    doc.add_heading('Signatures', level=1)
    
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.style = 'Table Grid'
    
    headers = sig_table.rows[0].cells
    headers[0].text = 'Party A'
    headers[1].text = 'Date'
    headers[2].text = 'Party B'
    
    sig_table.rows[1].cells[0].text = '___________________'
    sig_table.rows[1].cells[1].text = '___________________'
    sig_table.rows[1].cells[2].text = '___________________'
    
    sig_table.rows[2].cells[0].text = '{{PARTY_A_SIGNATURE}}'
    sig_table.rows[2].cells[1].text = '{{SIGNATURE_DATE}}'
    sig_table.rows[2].cells[2].text = '{{PARTY_B_SIGNATURE}}'
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/contract_template.docx')
    print('✓ Contract template created')


def create_memo_template():
    """Business memo template"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.add_run('MEMORANDUM').bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Memo fields
    memo_info = doc.add_paragraph()
    memo_info.add_run('TO: ').bold = True
    memo_info.add_run('{{TO}}\n')
    memo_info.add_run('FROM: ').bold = True
    memo_info.add_run('{{FROM}}\n')
    memo_info.add_run('DATE: ').bold = True
    memo_info.add_run('{{DATE}}\n')
    memo_info.add_run('SUBJECT: ').bold = True
    memo_info.add_run('{{SUBJECT}}')
    
    doc.add_paragraph()  # Spacing
    
    # Body
    doc.add_paragraph('{{MEMO_BODY}}')
    
    # Action Items (if any)
    doc.add_heading('Action Items', level=2)
    doc.add_paragraph('{{ACTION_ITEMS}}')
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/memo_template.docx')
    print('✓ Memo template created')


if __name__ == "__main__":
    print("Creating sample Word templates...\n")
    
    try:
        create_report_template()
        create_proposal_template()
        create_resume_template()
        create_email_template()
        create_contract_template()
        create_memo_template()
        
        print("\n" + "="*50)
        print("✓ All templates created successfully!")
        print("="*50)
        print("\nCreated templates:")
        print("  • templates/report_template.docx")
        print("  • templates/proposal_template.docx")
        print("  • templates/resume_template.docx")
        print("  • templates/email_template.docx")
        print("  • templates/contract_template.docx")
        print("  • templates/memo_template.docx")
        print("\nNext steps:")
        print("  1. Start backend: python -m uvicorn backend.main:app --reload")
        print("  2. Test endpoints in Swagger UI: http://localhost:8000/docs")
        print("  3. Use templates with /find-placeholders and /populate-template endpoints")
        
    except Exception as e:
        print(f"Error: {e}")
