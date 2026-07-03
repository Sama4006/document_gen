"""Template Engine - Handles Word document template population"""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import Dict, Optional
import os


class TemplateEngine:
    """Handle Word document (.docx) template population"""
    
    @staticmethod
    def find_placeholders(doc_path: str) -> list:
        """
        Find all placeholders in a Word document
        Placeholders should be in format: {{PLACEHOLDER_NAME}}
        
        Args:
            doc_path: Path to .docx template
            
        Returns:
            List of placeholder names found
        """
        doc = Document(doc_path)
        placeholders = set()
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if "{{" in text and "}}" in text:
                # Extract placeholder names
                import re
                matches = re.findall(r"{{(.*?)}}", text)
                placeholders.update(matches)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if "{{" in text and "}}" in text:
                            import re
                            matches = re.findall(r"{{(.*?)}}", text)
                            placeholders.update(matches)
        
        return sorted(list(placeholders))
    
    @staticmethod
    def populate_template(
        template_path: str,
        placeholders_data: Dict[str, str],
        output_path: Optional[str] = None
    ) -> str:
        """
        Populate a Word template with provided data
        
        Args:
            template_path: Path to .docx template
            placeholders_data: Dictionary of placeholder -> value mappings
            output_path: Output path for populated document
            
        Returns:
            Path to populated document
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        doc = Document(template_path)
        
        # Populate paragraphs
        for paragraph in doc.paragraphs:
            for key, value in placeholders_data.items():
                placeholder = "{{" + key + "}}"
                if placeholder in paragraph.text:
                    # Replace in runs to preserve formatting
                    TemplateEngine._replace_in_runs(paragraph, placeholder, str(value))
        
        # Populate tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in placeholders_data.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in paragraph.text:
                                TemplateEngine._replace_in_runs(paragraph, placeholder, str(value))
        
        # Save document
        if output_path is None:
            base, ext = os.path.splitext(template_path)
            output_path = f"{base}_populated{ext}"
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def _replace_in_runs(paragraph, placeholder: str, replacement: str):
        """
        Replace placeholder in paragraph runs while preserving formatting
        
        Args:
            paragraph: Paragraph object
            placeholder: Placeholder string to replace
            replacement: Replacement string
        """
        if placeholder in paragraph.text:
            # Merge runs and replace
            full_text = "".join([run.text for run in paragraph.runs])
            
            if placeholder in full_text:
                new_text = full_text.replace(placeholder, replacement)
                
                # Clear existing runs
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                # Add new run with replaced text
                new_run = paragraph.add_run(new_text)
                if paragraph.runs:
                    # Copy formatting from first run if available
                    first_run = paragraph.runs[0]
                    new_run.bold = first_run.bold
                    new_run.italic = first_run.italic
                    new_run.font.size = first_run.font.size


# Singleton instance
template_engine = TemplateEngine()
