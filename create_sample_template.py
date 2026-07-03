"""
Script to generate sample Word template with placeholders
Run once: python create_sample_template.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_sample_template():
    """Create a sample Word template with placeholders"""
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('{{TITLE}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run('Date: ').bold = True
    date_para.add_run('{{DATE}}')
    
    # Add executive summary section
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph('{{SUMMARY}}')
    
    # Add key points section
    doc.add_heading('Key Points', level=1)
    doc.add_paragraph('{{KEY_POINTS}}')
    
    # Add recommendations section
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('{{RECOMMENDATIONS}}')
    
    # Add table example
    doc.add_heading('Data Table', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Value'
    
    # Data rows
    table.rows[1].cells[0].text = 'Performance'
    table.rows[1].cells[1].text = '{{PERFORMANCE}}'
    
    table.rows[2].cells[0].text = 'Status'
    table.rows[2].cells[1].text = '{{STATUS}}'
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('{{CONCLUSION}}')
    
    # Create templates directory if needed
    os.makedirs('templates', exist_ok=True)
    
    # Save template
    output_path = 'templates/sample_template.docx'
    doc.save(output_path)
    print(f"✓ Sample template created: {output_path}")
    print(f"\nPlaceholders found:")
    print("  - {{TITLE}}")
    print("  - {{DATE}}")
    print("  - {{SUMMARY}}")
    print("  - {{KEY_POINTS}}")
    print("  - {{RECOMMENDATIONS}}")
    print("  - {{PERFORMANCE}}")
    print("  - {{STATUS}}")
    print("  - {{CONCLUSION}}")


if __name__ == "__main__":
    create_sample_template()
